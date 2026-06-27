# -*- coding: utf-8 -*-
"""產生決賽顧問諮詢提問表 .docx（對齊官方三大重點：Rehearsal／Model 調校／平台障礙）。

定稿原則：顧問主時間用來壓測故事、Demo 節奏、商業價值；
工程規格題會前先實測消掉，未解者才帶截圖以平台障礙形式問。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "Microsoft JhengHei"
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0xB0, 0x4A, 0x1F)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = CJK
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.2)


def set_cjk(run, name=CJK):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def para(text="", size=11, bold=False, color=None, align=None,
         space_after=6, space_before=0, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        set_cjk(r)
    return p


def section(title):
    p = para(title, size=13, bold=True, color=NAVY, space_before=10, space_after=5)
    p.paragraph_format.keep_with_next = True
    return p


def question(num, topic, q, rationale):
    """單題：粗體編號＋主題，問題本文，灰字『想釐清』一行。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{num}. {topic}")
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY
    set_cjk(r)

    pb = doc.add_paragraph()
    pb.paragraph_format.space_after = Pt(2)
    pb.paragraph_format.left_indent = Cm(1.0)
    pb.paragraph_format.keep_with_next = True
    rb = pb.add_run(q)
    rb.font.size = Pt(11)
    set_cjk(rb)

    pr = doc.add_paragraph()
    pr.paragraph_format.space_after = Pt(8)
    pr.paragraph_format.left_indent = Cm(1.0)
    rr = pr.add_run("想釐清：" + rationale)
    rr.font.size = Pt(9.5)
    rr.font.color.rgb = GREY
    set_cjk(rr)


def compact_question(num, topic, q):
    """附件題：保持完整問句，但用較緊湊的版面呈現。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f"{num}. {topic}：")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    set_cjk(r)
    body = p.add_run(q)
    body.font.size = Pt(10.5)
    set_cjk(body)


# ── 標題 ─────────────────────────────────────────
para("2026 精誠 AI 創新競賽 — 決賽顧問諮詢提問表", size=16, bold=True,
     color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("隊伍序號：32　　隊伍名稱：Safari 叫爸爸　　提案名稱：ETF 透視鏡 PASSPORT",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
para("提交日期：2026/06/05", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=GREY, space_after=8)

# ── 一、專案簡述 ─────────────────────────────────
section("一、專案簡述")
para(
    "ETF 透視鏡 PASSPORT 以「散戶 30 秒看穿自己手上多檔 ETF 的穿透曝險」為核心。"
    "使用者輸入自己的持股組合與金額，系統以 Hybrid RAG（Graph＋Vector）穿透至個股層，"
    "計算如「台積電在你整體組合的穿透占比與新台幣金額」，並輔以穿透流向圖、表格與引用回鏈。"
    "核心論點：散戶以為買多檔 ETF 是分散，其實可能高度重複曝險於同一檔個股。"
    "資料涵蓋 16 檔台股 ETF（12 檔散戶熱門＋4 檔主動式對照組），持股為真實成分股、快照日 2026-06-05，"
    "前端透過 EAP Chat API 串接，現場 Demo 為決賽主要展示。"
    "持股權重來源 etfinfo.tw，經元大投信官網直抓交叉驗證，評審可拿公開持股當場核對；"
    "質化題（如收益平準金）由 EAP Vector RAG 檢索我方已上傳的 11 份真實 SITCA 配息公告佐證。",
    size=11, space_after=6)

# ── 二、核心展示流程草案 ─────────────────────────
section("二、核心展示流程草案（請協助檢視）")
para("我們預計將決賽核心展示控制在 90 秒內。若諮詢時版本已備妥，希望以畫面快速走過一次；"
     "若尚在整合，也盼顧問先就下方流程協助檢視敘事與取捨。",
     size=11, space_after=4)
para("預計 90 秒展示流程：", size=11, bold=True, color=ACCENT, space_after=2,
     indent=0.6)
para("0050 ＋ 006208 ＋ 00878 ＋ 00992A 各投 10 萬 → 穿透試算出台積電曝險 31.1%"
     "→ 穿透流向圖呈現「4 種包裝、同一籃台積電」→ 點開引用回鏈驗證來源。",
     size=11, space_after=6, indent=1.0)

# ── 三、優先確認 5 題 ─────────────────────────────
section("三、優先確認（5 題，顧問主時間）")
question(
    1, "核心展示情境的風險檢視",
    "如果您是評審，會從哪個角度不相信這個 31.1%？資料快照、金融合規、"
    "Hybrid 必要性、還是計算可信度，哪一個最危險？",
    "希望辨識最需要補強的風險，作為決賽 Q&A 的準備核心。")
question(
    2, "後端計算架構會不會被當成「只是試算器」",
    "因實測 EAP 自動產 Cypher 不可靠（ticker 存成 float、LLM 重寫並丟參數），我們把穿透／排名等"
    "數字改由後端 deterministic 計算（penetrate.py 讀同一份快照），EAP 只負責質化敘述與引用。"
    "這個取捨評審會不會質疑「沒真的用到 Graph RAG，只是 Python 試算器＋Chatbot」？"
    "該如何在 Demo 與簡報證明 Graph／Hybrid 的實質價值？",
    "把工程現實（EAP 自動 Cypher 不穩）與評分（技術應用深度 20%）對齊的關鍵取捨。")
question(
    3, "Demo 節奏與技術深度",
    "決賽 5 分鐘內，我們應該主打一個核心查詢情境，還是展示 Graph／Vector／Hybrid "
    "三題以證明完整性？技術說明應佔多少比例？",
    "同時決定 Demo 30%、技術 20%、介面 15% 的平衡；技術講太多會失焦、太少又怕不像 Hybrid RAG。")
question(
    4, "使用情境與商業落地",
    "決賽 Demo 是否適合以散戶作為使用情境，讓痛點一秒可懂；商業落地則主打券商理專輔助"
    "或投信客服／投資教育？哪一個付費方與導入路徑最合理？",
    "區分實際使用者與付費客戶，讓商業模式更具體。")
question(
    5, "自建前端如何證明 EAP 技術深度",
    "我們預計以自建前端呈現穿透流向圖與引用回鏈。決賽時應如何證明底層確實使用 EAP Hybrid RAG，"
    "而非一般試算器或 Chatbot？是否需要展示平台後台、查詢路徑或 Graph／Vector 分流結果？",
    "希望兼顧前端體驗與技術應用深度，避免展示焦點失衡。")

# ── 四、效果調校與資料模型 ───────────────────────
section("四、效果調校與資料模型（時間允許時確認）")
question(
    6, "Model 約束與引用可信度",
    "Robot Setting 是否足以約束「數字一律走 Graph、不准 LLM 心算」？是否建議改為固定 "
    "Cypher template ＋ LLM 只填參數？另若 trace（cyphers／rag_chunk_info）拿不到，"
    "前端用 Graph 查詢結果＋文件 metadata 自行組 citation，是否可被接受？",
    "把 Model 效果調校與引用可信度合併，呼應官方第二點。")
question(
    7, "Q-D 名實／Q-E 漂綠要不要列入主秀",
    "漂綠／名實檢測需要個股 ESG 評分與主題標籤；但換成真實成分股後，我們的示意 ESG／主題只覆蓋"
    "部分個股（如 00878 可評估權重僅約 39%），可信度下降。決賽是否該聚焦 Q-C 穿透＋Q-A 排名＋"
    "Q-B 質化（皆為真實資料），把 Q-D／Q-E 當延伸帶過？還是值得投資補真實 ESG（如 TWSE 公司治理"
    "評鑑 t187ap46）讓它們站得住、敢被評審追問？",
    "避免主打一個資料覆蓋不足、評審一問就破的延伸題。")

# ── 五、決賽現場與備援規則 ───────────────────────
doc.add_page_break()
section("五、決賽現場與備援規則")
para("以下為現場安排相關問題，盼能提早準備展示設備與備援流程。",
     size=10.5, color=GREY, space_after=4)
compact_question(8, "報告規則、評分權重與送件",
                 "決賽簡報是否有頁數限制？簡報、現場 Demo 與 Q&A 各分配多少時間？"
                 "決賽評分各構面權重（Demo 展示 30%／技術應用 20%／介面 15%／…）是否即為最終比重？"
                 "另：決賽簡報上傳表單連結與截止日（6/21 23:59？）為何？")
compact_question(9, "現場設備",
                 "現場 Demo 使用參賽團隊自備電腦，還是由主辦方統一提供設備？")
compact_question(10, "Demo 備援",
                 "若現場網路或 API 不穩，是否接受本地備援流程或事先錄製影片？建議如何說明以兼顧評分公平性？")

# ── 六、平台問題回報（已實測，附重現） ──────────────
section("六、平台問題回報（已實測，請工程團隊協助）")
para("以下為會前實測 EAP 時確認的平台匯入／設定問題，附重現方式，懇請工程團隊協助。",
     size=10.5, color=GREY, space_after=4)
compact_question(
    11, "ticker 欄位被當數值匯入，前導零消失、型別不一致",
    "CSV 內如 0050、00878 等純數字 ETF／個股代號，匯入後在 graph 被存成浮點數："
    "toString(ticker) 回傳「2330.0」、0050 變成 50，前導零消失；而帶字母的主動式代號"
    "（如 00992A）則正常存為字串，造成同一 ticker 欄位型別不一致。後果是 LLM 生成的 "
    "toString(ticker)='0050' 比對全部落空，需用 toInteger() 才勉強匹配。"
    "請問匯入時能否強制指定 ticker 欄位以字串型別處理、保留前導零？")
compact_question(
    12, "Robot Setting／System Prompt 找不到可用設定入口",
    "我們需要把最新 robot_setting.txt 貼到 EAP assistant，讓質化題穩定 grounding、避免 LLM 憑記憶回答。"
    "（實測佐證：知識庫已成功上傳並索引 00878 配息公告，但泛問『00878 最近一次配息平準金占多少』時，"
    "LLM 仍回訓練資料的假數字 78/12/10；改問『依 112/10/31 公告』才正確回 5.71%。代表單純上傳文件不夠，"
    "需要設定強制『檢索優先、不准憑記憶』。）"
    "已實測：Assistant 頁會自動進入 /portal/project/{project_id}/assistant/chat/{chat_id}，"
    "前端 bundle 只看到 assistant chat/shared 與一般 Project Settings，未看到機器人設定或 system prompt route；"
    "Portal GraphQL introspection 只列出 project/chat/message/question 等 mutation，ProjectEditInput 也沒有 system prompt/"
    "assistant setting 欄位；/api/v1/robot/setting 對 POST 回 405 Not Allowed。"
    "請問目前 EAP 要在哪裡設定 Robot Setting？是否需要特定權限、隱藏入口，或由工程協助開啟／匯入？")

# 會前已實測釐清、無需提問：chat body field = q、streaming 為 SSE（data: 行）、
# Robot Setting 需新開 chat 生效、cyphers 欄位確實有回傳；但目前 UI/API 設定入口待工程確認（Q12）。
# Vector 知識庫上傳已自行解掉：POST /import/vector/knowledge 要用 multipart/form-data（JSON 會 502），
#   11 份 SITCA 配息公告已上傳並索引完成，故不列為提問。
# 另一觀察（已併入 Q6）：EAP chat 的 LLM 會自行重寫指定的 Cypher、且可能丟失參數，
# 故我方 Demo 的數字一律由後端自算，EAP 僅負責質化敘述。

para("感謝諮詢委員撥冗指導。", size=11, color=GREY, space_before=10)

out = "32_Safari叫爸爸_諮詢提問表.docx"
doc.save(out)
print("saved", out)
